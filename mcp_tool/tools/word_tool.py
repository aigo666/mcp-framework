"""
Word文档解析工具，用于解析Word文档内容
"""

import os
import traceback
import io
import base64
import imghdr
import logging
from typing import Dict, List, Any, Tuple
import docx
from docx.document import Document
from docx.parts.document import DocumentPart
from docx.package import Package
import mcp.types as types
from . import BaseTool, ToolRegistry

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@ToolRegistry.register
class WordTool(BaseTool):
    """
    用于解析Word文档的工具，提取文本内容、表格和图片信息
    """
    
    name = "parse_word"
    description = "解析Word文档内容，提取文本、表格和图片信息"
    input_schema = {
        "type": "object",
        "required": ["file_path"],
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Word文档的本地路径，例如'/path/to/document.docx'",
            }
        },
    }
    
    async def execute(self, arguments: Dict[str, Any]) -> List[types.TextContent | types.ImageContent | types.EmbeddedResource]:
        """
        解析Word文档
        
        Args:
            arguments: 参数字典，必须包含'file_path'键
            
        Returns:
            解析结果列表
        """
        logger.info(f"开始处理Word文档: {arguments.get('file_path', 'Unknown')}")
        
        if "file_path" not in arguments:
            return [types.TextContent(
                type="text",
                text="错误: 缺少必要参数 'file_path'"
            )]
        
        # 处理文件路径，支持挂载目录的转换
        file_path = self.process_file_path(arguments["file_path"])
        logger.info(f"处理后的文件路径: {file_path}")
        
        return await self._parse_word_document(file_path)
    
    def _get_image_mime_type(self, image_bytes: bytes) -> str:
        """
        获取图片的MIME类型
        """
        try:
            image_type = imghdr.what(None, image_bytes)
            if image_type:
                return f"image/{image_type}"
            return "image/png"  # 默认返回PNG类型
        except Exception as e:
            logger.error(f"获取图片MIME类型失败: {str(e)}")
            return "image/png"  # 出错时默认返回PNG类型
    
    def _encode_image_base64(self, image_bytes: bytes) -> str:
        """
        将图片编码为base64格式
        """
        try:
            return base64.b64encode(image_bytes).decode('utf-8')
        except Exception as e:
            logger.error(f"Base64编码失败: {str(e)}")
            return ""
    
    def _extract_images_from_word(self, doc: Document) -> List[Tuple[str, bytes]]:
        """
        从Word文档中提取图片
        
        Args:
            doc: Word文档对象
            
        Returns:
            图片列表，每项包含图片ID和二进制数据
        """
        images = []
        try:
            document_part = doc.part
            rels = document_part.rels
            
            logger.info(f"文档关系数量: {len(rels)}")
            
            for rel_id, rel in rels.items():
                try:
                    # 更精确的图片类型判断
                    is_image = False
                    if hasattr(rel, 'reltype') and rel.reltype:
                        is_image = 'image' in rel.reltype.lower()
                    elif hasattr(rel, '_target') and rel._target:
                        is_image = any(ext in rel._target.lower() for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'])
                    
                    # 输出关系信息用于调试
                    logger.info(f"关系ID: {rel_id}, 类型: {getattr(rel, 'reltype', 'Unknown')}, 目标: {getattr(rel, '_target', 'Unknown')}")
                    
                    if is_image:
                        try:
                            image_part = rel.target_part
                            if hasattr(image_part, 'blob') and image_part.blob:
                                image_bytes = image_part.blob
                                logger.info(f"成功提取图片 {rel_id}, 大小: {len(image_bytes)} 字节")
                                images.append((rel_id, image_bytes))
                            else:
                                logger.warning(f"图片部分没有blob属性或blob为空: {rel_id}")
                        except Exception as part_error:
                            logger.error(f"处理图片部分时出错: {str(part_error)}")
                except Exception as rel_error:
                    logger.error(f"处理关系时出错: {str(rel_error)}")
            
            # 尝试使用替代方法提取图片
            if not images:
                logger.info("尝试使用替代方法提取图片...")
                try:
                    # 递归查找所有内联图片
                    for shape in doc.inline_shapes:
                        if shape._inline.graphic:
                            try:
                                blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                                rId = blip.embed
                                image_part = doc.part.related_parts[rId]
                                if hasattr(image_part, 'blob') and image_part.blob:
                                    image_bytes = image_part.blob
                                    logger.info(f"从内联图形中提取图片 {rId}, 大小: {len(image_bytes)} 字节")
                                    images.append((rId, image_bytes))
                            except Exception as shape_error:
                                logger.error(f"处理内联图形时出错: {str(shape_error)}")
                except Exception as shapes_error:
                    logger.error(f"处理所有内联图形时出错: {str(shapes_error)}")
        except Exception as e:
            logger.error(f"提取图片过程中出错: {str(e)}")
        
        logger.info(f"总共提取到 {len(images)} 张图片")
        return images
    
    async def _parse_word_document(self, file_path: str) -> List[types.TextContent | types.ImageContent | types.EmbeddedResource]:
        """
        解析Word文档内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            Word文档内容列表
        """
        results = []
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return [types.TextContent(
                type="text",
                text=f"错误: 文件不存在: {file_path}\n请检查路径是否正确，并确保文件可访问。"
            )]
        
        # 检查文件扩展名
        if not file_path.lower().endswith(('.docx', '.doc')):
            logger.error(f"不支持的文件格式: {file_path}")
            return [types.TextContent(
                type="text",
                text=f"错误: 不支持的文件格式: {file_path}\n仅支持.docx和.doc格式的Word文档。"
            )]
        
        try:
            # 添加文件信息
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            logger.info(f"文件大小: {file_size_mb:.2f} MB")
            results.append(types.TextContent(
                type="text",
                text=f"# Word文档解析\n\n文件名: {os.path.basename(file_path)}\n文件大小: {file_size_mb:.2f} MB"
            ))
            
            # 打开Word文档
            logger.info(f"开始打开Word文档: {file_path}")
            doc = docx.Document(file_path)
            logger.info("Word文档打开成功")
            
            # 提取文档属性
            properties = {}
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                properties['标题'] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                properties['作者'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                properties['创建时间'] = str(doc.core_properties.created)
            if hasattr(doc.core_properties, 'modified') and doc.core_properties.modified:
                properties['修改时间'] = str(doc.core_properties.modified)
            if hasattr(doc.core_properties, 'comments') and doc.core_properties.comments:
                properties['备注'] = doc.core_properties.comments
            
            # 添加文档属性信息
            if properties:
                properties_text = "## 文档属性\n\n"
                for key, value in properties.items():
                    properties_text += f"- {key}: {value}\n"
                results.append(types.TextContent(
                    type="text",
                    text=properties_text
                ))
            
            # 提取文档内容
            content_text = "## 文档内容\n\n"
            
            # 处理段落
            paragraphs_count = len(doc.paragraphs)
            logger.info(f"文档包含 {paragraphs_count} 个段落")
            content_text += f"### 段落 (共{paragraphs_count}个)\n\n"
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():  # 只处理非空段落
                    content_text += f"{para.text}\n\n"
            
            # 处理表格
            tables_count = len(doc.tables)
            if tables_count > 0:
                logger.info(f"文档包含 {tables_count} 个表格")
                content_text += f"### 表格 (共{tables_count}个)\n\n"
                
                for i, table in enumerate(doc.tables):
                    content_text += f"#### 表格 {i+1}\n\n"
                    
                    # 创建Markdown表格
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                        rows.append(cells)
                    
                    if rows:
                        # 表头
                        content_text += "| " + " | ".join(rows[0]) + " |\n"
                        # 分隔线
                        content_text += "| " + " | ".join(["---"] * len(rows[0])) + " |\n"
                        # 表格内容
                        for row in rows[1:]:
                            content_text += "| " + " | ".join(row) + " |\n"
                        
                        content_text += "\n"
            
            # 添加文档内容
            results.append(types.TextContent(
                type="text",
                text=content_text
            ))
            
            # 提取图片信息和内容
            try:
                logger.info("开始提取文档中的图片")
                # 提取文档中的所有图片
                images = self._extract_images_from_word(doc)
                
                if images:
                    image_info = f"## 图片信息\n\n文档中包含 {len(images)} 张图片。\n\n"
                    results.append(types.TextContent(
                        type="text",
                        text=image_info
                    ))
                    
                    # 返回图片内容
                    for i, (image_id, image_bytes) in enumerate(images):
                        try:
                            # 获取图片MIME类型
                            mime_type = self._get_image_mime_type(image_bytes)
                            logger.info(f"图片 {i+1} MIME类型: {mime_type}")
                            
                            # 将图片添加到结果中
                            image_base64 = self._encode_image_base64(image_bytes)
                            if image_base64:
                                results.append(types.TextContent(
                                    type="text",
                                    text=f"### 图片 {i+1}\n\n"
                                ))
                                
                                # 将图片作为ImageContent类型添加
                                logger.info(f"添加图片 {i+1} 到结果中")
                                results.append(types.ImageContent(
                                    type="image",
                                    data=image_base64,
                                    mimeType=mime_type
                                ))
                            else:
                                logger.warning(f"图片 {i+1} Base64编码为空")
                        except Exception as img_error:
                            error_details = traceback.format_exc()
                            logger.error(f"处理图片 {i+1} 时出错: {str(img_error)}\n{error_details}")
                            results.append(types.TextContent(
                                type="text",
                                text=f"警告: 处理图片 {i+1} 时出错: {str(img_error)}"
                            ))
                else:
                    logger.info("文档中未包含图片")
                    results.append(types.TextContent(
                        type="text",
                        text="## 图片信息\n\n文档中未包含图片。"
                    ))
            except Exception as img_error:
                error_details = traceback.format_exc()
                logger.error(f"提取图片信息时出错: {str(img_error)}\n{error_details}")
                results.append(types.TextContent(
                    type="text",
                    text=f"警告: 提取图片信息时出错: {str(img_error)}"
                ))
            
            # 添加处理完成的提示
            logger.info("Word文档处理完成")
            results.append(types.TextContent(
                type="text",
                text="Word文档处理完成！"
            ))
            
            return results
        except Exception as e:
            error_details = traceback.format_exc()
            logger.error(f"解析Word文档失败: {str(e)}\n{error_details}")
            return [types.TextContent(
                type="text",
                text=f"错误: 解析Word文档失败: {str(e)}\n"
                     f"可能的原因:\n"
                     f"1. 文件格式不兼容或已损坏\n"
                     f"2. 文件受密码保护\n"
                     f"3. 文件包含不支持的内容\n\n"
                     f"详细错误信息: {error_details}"
            )] 